# File-based Q&A Retrieval with source Chatbot using LangChain, OpenAI, Qdrant, HuggingFace, and Streamlit
# Upload PDFs, embed chunks, store in Qdrant, retrieve via similarity search, chat with your documents.
import streamlit as st
import os
from PyPDF2 import PdfReader
import docx
from langchain.chat_models import ChatOpenAI
from langchain.vectorstores import Qdrant
import random
from datetime import datetime
from langchain import PromptTemplate
from langchain.chains import RetrievalQA
import string
from qdrant_client import QdrantClient
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from streamlit_chat import message
from langchain.callbacks import get_openai_callback
from langchain.docstore.document import Document
openapi_key = st.secrets["OPENAI_API_KEY"]
qdrant_url = st.secrets["QDRANT_URL"]
qdrant_api_key = st.secrets["QDRANT_API_KEY"]
embeddings = HuggingFaceEmbeddings(model_name = "sentence-transformers/all-mpnet-base-v2")
def main():
    load_dotenv()
    st.set_page_config(page_title = "Q/A with your file")
    st.header("Retrieval QA Chain")
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "processComplete" not in st.session_state:
        st.session_state.processComplete = None
    with st.sidebar:
        uploaded_files = st.file_uploader("Upload your file", type = ["pdf"], accept_multiple_files = True)
        openai_api_key = openapi_key
        openai_api_key = st.text_input("OpenAI API Key", key = openapi_key, type = "password")
        process = st.button("Process")
    if process:
        if not openai_api_key:
            st.info("Please add your OpenAI API Key to continue.")
            st.stop()
        text_chunks_list = []
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_text = get_files_text(uploaded_file)
            text_chunks = get_text_chunks(file_text, file_name)
            text_chunks_list.extend(text_chunks)
        curr_date = str(datetime.now())
        collection_name = "".join(random.choices(string.ascii_letters, k = 4)) + curr_date.split(".")[0].replace(":", "-").replace(" ", "T")
        vectorstore = get_vectorstore(text_chunks_list, collection_name)
        st.write("Vector Store Created...")
        num_chunks = 4
        st.session_state.conversation = get_qa_chain(vectorstore, num_chunks)
        st.session_state.processComplete = True
    if st.session_state.processComplete == True:
        user_question = st.chat_input("Ask Question about your files.")
        if user_question:
            handle_userinput(user_question)


def get_files_text(uploaded_file):
    text = ""
    split_up = os.path.splitext(uploaded_file.name)
    file_extension = split_up[1]
    if file_extension == ".pdf":
        text += get_pdf_text(uploaded_file)
    elif file_extension == ".docx":
        text += get_docx_text(uploaded_file)
    else:
        pass
    return text


def get_pdf_text(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def get_docx_text(file):
    doc = docx.Document(file)
    allText = []
    for docpara in doc.paragraphs:
        allText.append(docpara.text)
    text = ' '.join(allText)
    return text

def get_text_chunks(text, filename):
    text_splitter = CharacterTextSplitter(
        separator = "\n",
        chunk_size = 80,
        chunk_overlap = 20,
        length_function = len
    )
    chunks = text_splitter.split_text(text)
    doc_list = []
    for chunk in chunks:
        metadata = {"source": filename}
        doc_string = Document(page_content = chunk, metadata = metadata)
        doc_list.append(doc_string)
    return doc_list


def get_vectorstore(text_chunks, COLLECTION_NAME):
    try:
        knowledge_base = Qdrant.from_documents(
            documents = text_chunks,
            embedding = embeddings,
            url = qdrant_url,
            prefer_grpc = True,
            api_key = qdrant_api_key,
            collection_name = COLLECTION_NAME
        )
    except Exception as e:
        st.write(f"Error: {e}")
    return knowledge_base


def get_qa_chain(vectorstore, num_chunks):
    qa = RetrievalQA.from_chain_type(
        llm = ChatOpenAI(
            model = "gpt-3.5-turbo"
        ),
        chain_type = "stuff",
        retriever = vectorstore.as_retriever(
            search_type = "similarity",
            search_kwargs = {"k": num_chunks}
        ),
        return_source_documents = True
    )
    return qa


def handle_userinput(user_question):
    with st.spinner("Generating response..."):
        result = st.session_state.conversation({"query": user_question})
        response = result["result"]
        source = result['source_documents'][0].metadata['source']
    st.session_state.chat_history.append(user_question)
    st.session_state.chat_history.append(f"{response} \nSource Document: {source}")
    response_container = st.container()
    with response_container:
        for i, messages in  enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                message(messages, is_user = True, key = str(i))
            else:
                message(messages, key = str(i))


if __name__ == "__main__":
    main()
